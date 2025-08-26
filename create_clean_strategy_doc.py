#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def add_text(doc, text, bold=False, size=11, indent=0):
    """添加文本段落，不使用项目符号"""
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent * 0.3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def create_strategy_document():
    # 创建文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加标题
    title = doc.add_heading('20万资金两阶段炒股策略', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph('基于动量效应的量化投资框架')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # ========== 总体框架 ==========
    doc.add_heading('一、总体投资框架', 1)
    
    doc.add_heading('（一）投资目标', 2)
    add_text(doc, '阶段一：筹码倍增期（6-12个月）——追求资产快速增值')
    add_text(doc, '阶段二：稳健复利期（长期框架）——追求持续稳定收益')
    
    doc.add_heading('（二）资金规模与时间安排', 2)
    add_text(doc, '初始资金：20万人民币', bold=True)
    add_text(doc, '冲刺窗口：6-18个月（默认9-12个月）', bold=True)
    add_text(doc, '切换条件：盈利达标或风险触发自动切换', bold=True)
    
    doc.add_page_break()
    
    # ========== 阶段一详细策略 ==========
    doc.add_heading('二、阶段一：筹码倍增策略详解', 1)
    
    doc.add_heading('（一）资产配置结构', 2)
    
    # 1. 核心持仓部分
    doc.add_heading('1. 核心持仓（60%）——长期底座配置', 3)
    
    # 创建核心持仓表格
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '资产类别'
    header_cells[1].text = '配置比例'
    header_cells[2].text = '具体标的'
    header_cells[3].text = '配置理由'
    
    # 设置表头格式
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 填充数据
    data = [
        ['A股宽基', '20%', '沪深300 ETF', '市场基准，低波动'],
        ['红利因子', '10%', '上证红利/中证红利低波ETF', '稳定分红，防守属性'],
        ['短期固收', '15%', '短融/短债ETF', '流动性管理，低风险'],
        ['黄金', '10%', '黄金ETF', '避险配置，对冲通胀'],
        ['美股基线', '5%', '标普500 QDII-ETF', '全球配置（仅溢价≤3%时买入）']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # 动态调整规则
    add_text(doc, '动态调整规则：', bold=True, size=12)
    add_text(doc, '1. 当10年国债收益率≥2.20%时：短融向5-10年国债ETF转移5-10%', indent=1)
    add_text(doc, '2. 当10年国债收益率≤1.60%时：减持中长债5%回短融', indent=1)
    add_text(doc, '3. 黄金价格>840元/克时：回落到目标10%；<750元/克时：补回到10%', indent=1)
    
    doc.add_paragraph()
    
    # 2. 卫星持仓部分
    doc.add_heading('2. 卫星持仓（40%）——进攻配置', 3)
    
    doc.add_heading('2.1 行业/主题动量轮动（20%）', 4)
    
    add_text(doc, '候选ETF池：', bold=True, size=12)
    add_text(doc, '科技类：半导体ETF、计算机ETF、科创50', indent=1)
    add_text(doc, '新经济：新能源车ETF、医药ETF', indent=1)
    add_text(doc, '周期类：有色ETF、银行ETF、证券ETF', indent=1)
    add_text(doc, '主题类：军工ETF', indent=1)
    
    doc.add_paragraph()
    add_text(doc, '轮动策略核心公式：', bold=True, size=12)
    add_text(doc, '评分 = 0.6 × 近3个月收益率 + 0.4 × 近6个月收益率', indent=1)
    add_text(doc, '持仓规则：持有评分前2名，各配置10%（单只上限12%）', indent=1)
    
    doc.add_paragraph()
    add_text(doc, '风险控制机制：', bold=True, size=12)
    
    # 创建风控表格
    risk_table = doc.add_table(rows=7, cols=2)
    risk_table.style = 'Table Grid'
    
    risk_header = risk_table.rows[0].cells
    risk_header[0].text = '控制维度'
    risk_header[1].text = '具体规则'
    
    for cell in risk_header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    risk_data = [
        ['缓冲带设置', '默认3%（强趋势降至2%，震荡市场升至4%）'],
        ['最短持有期', '≥2周（止损情况除外）'],
        ['趋势闸门', '沪深300在年线下方：仅持1只×15%\n沪深300在年线上方：持2只×各10%'],
        ['相关性控制', '前二名日相关系数>0.8时，用第三名替换第二名'],
        ['止损线', '默认-12%（强趋势-10%，震荡市-15%）'],
        ['锁盈线', '任一持仓+25%时，抽出盈利50%并入核心']
    ]
    
    for i, row_data in enumerate(risk_data, 1):
        row_cells = risk_table.rows[i].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
    
    doc.add_page_break()
    
    # 2.2 可转债网格交易
    doc.add_heading('2.2 可转债网格交易（10%）', 4)
    
    add_text(doc, '入池标准：', bold=True, size=12)
    add_text(doc, '基本要求：规模≥5亿，评级≥AA-，剩余期限1-5年', indent=1)
    add_text(doc, '价格区间：90-130元', indent=1)
    add_text(doc, '转股溢价：0-40%（高位市场可放宽至45%）', indent=1)
    add_text(doc, '流动性要求：近3个月日均成交额≥2000万', indent=1)
    
    doc.add_paragraph()
    add_text(doc, '网格设置方法：', bold=True, size=12)
    add_text(doc, '动态间距 = max(3%, 0.8 × 20日ATR/价格)', indent=1)
    add_text(doc, '间距限制：2-6%', indent=1)
    add_text(doc, '可采用双层结构：内层±2-3%，外层±4-5%', indent=1)
    
    doc.add_paragraph()
    add_text(doc, '仓位管理规则：', bold=True, size=12)
    add_text(doc, '单券上限：1.5%', indent=1)
    add_text(doc, '总仓上限：10%', indent=1)
    add_text(doc, '分散度要求：8-10只', indent=1)
    
    # 可转债评分系统表格
    doc.add_paragraph()
    add_text(doc, '可转债评分系统（满分1.0）：', bold=True, size=12)
    
    score_table = doc.add_table(rows=10, cols=3)
    score_table.style = 'Table Grid'
    
    score_header = score_table.rows[0].cells
    score_header[0].text = '因子类型'
    score_header[1].text = '权重'
    score_header[2].text = '评价维度'
    
    for cell in score_header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    score_data = [
        ['正向因子合计', '0.95', ''],
        ['价值垫', '0.22', '距离债底的安全边际'],
        ['转股溢价', '0.18', '溢价率越低越好'],
        ['流动性', '0.18', '成交额(70%)，买卖价差(30%)'],
        ['可交易波动', '0.12', '日波动2%/日最佳'],
        ['期限匹配', '0.10', '2.2年最佳，1.5-3年较优'],
        ['信用质量', '0.10', 'AAA=1.0, AA+=0.9, AA=0.8, AA-=0.7'],
        ['条款友好', '0.05', '下修条款(60%)，回售条款(40%)'],
        ['负向因子合计', '-0.15', '强赎风险(-0.10)，事件风险(-0.05)']
    ]
    
    for i, row_data in enumerate(score_data, 1):
        row_cells = score_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    add_text(doc, '入围标准：综合得分≥0.60', bold=True)
    
    doc.add_paragraph()
    doc.add_heading('2.3 战术/机会仓位（10%）', 4)
    add_text(doc, '1. 无合适机会时停放货币基金')
    add_text(doc, '2. 用于把握突发性投资机会')
    add_text(doc, '3. 严格控制风险，快进快出')
    
    doc.add_page_break()
    
    # ========== 执行方案 ==========
    doc.add_heading('三、执行方案', 1)
    
    doc.add_heading('（一）建仓计划（6-8周DCA）', 2)
    add_text(doc, '6周方案：每周投入16.7%')
    add_text(doc, '8周方案：每周投入12.5%')
    add_text(doc, '执行日：每周二/周五')
    add_text(doc, '下单时间：10:30或14:00')
    add_text(doc, '执行方式：限价单，偏离IOPV≤0.1%')
    
    doc.add_heading('（二）再平衡规则', 2)
    
    add_text(doc, '季度观察：', bold=True)
    add_text(doc, '触发条件：任一大类偏离目标±5个百分点', indent=1)
    add_text(doc, '调整幅度：回调至±2个百分点内', indent=1)
    
    add_text(doc, '年度再平衡：', bold=True)
    add_text(doc, '时间：每年12月最后一周', indent=1)
    add_text(doc, '动作：所有资产回到目标权重', indent=1)
    
    add_text(doc, '卫星轮动：', bold=True)
    add_text(doc, '频率：每月最后一个交易日', indent=1)
    add_text(doc, '方法：重新计算动量得分并调整持仓', indent=1)
    
    doc.add_heading('（三）组合级风控', 2)
    
    # 风控表格
    drawdown_table = doc.add_table(rows=4, cols=2)
    drawdown_table.style = 'Table Grid'
    
    drawdown_header = drawdown_table.rows[0].cells
    drawdown_header[0].text = '回撤幅度'
    drawdown_header[1].text = '调整措施'
    
    for cell in drawdown_header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    drawdown_data = [
        ['-15%', '卫星从40%降至25%（行业-10%，转债-5%）'],
        ['-20%', '卫星降至20%（行业12%，转债8%）'],
        ['-30%', '卫星降至10%（行业≤5%，转债≤5%），其余回核心/现金']
    ]
    
    for i, row_data in enumerate(drawdown_data, 1):
        row_cells = drawdown_table.rows[i].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
    
    doc.add_page_break()
    
    # ========== 阶段切换条件 ==========
    doc.add_heading('四、阶段切换条件', 1)
    
    switch_table = doc.add_table(rows=5, cols=3)
    switch_table.style = 'Table Grid'
    
    switch_header = switch_table.rows[0].cells
    switch_header[0].text = '触发类型'
    switch_header[1].text = '具体条件'
    switch_header[2].text = '立即行动'
    
    for cell in switch_header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    switch_data = [
        ['盈利达标', 
         '1.冲刺仓收益≥100%（6万→12万）\n2.总资产≥30万', 
         '冲刺盈利2/3并入核心\n冲刺仓降至20%'],
        ['风险止损', 
         '1.冲刺仓峰值回撤≥30%\n2.总资产≤14万', 
         '终止冲刺\n切换长期框架'],
        ['时限到期', 
         '满12个月未达标\n且Calmar<0.5', 
         '结束冲刺\n转入长期框架'],
        ['绩效体检', 
         '滚动6个月Calmar<0.5\n连续两次', 
         '冲刺仓从30%降至20%']
    ]
    
    for i, row_data in enumerate(switch_data, 1):
        row_cells = switch_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # ========== 策略科学依据 ==========
    doc.add_heading('五、策略科学依据', 1)
    
    doc.add_heading('（一）动量效应理论支撑', 2)
    add_text(doc, '1. 经典研究：买入近3-12个月赢家，卖出输家，长期有显著正超额收益')
    add_text(doc, '2. 跨市场验证：在8大资产类别、多个国家长期存在')
    add_text(doc, '3. 行业层面更稳健：个股动量很大部分来自行业动量')
    
    doc.add_heading('（二）A股市场特征适配', 2)
    add_text(doc, '1. 行业轮动明显：A股板块效应强，适合行业ETF轮动')
    add_text(doc, '2. 个股动量弱化：2005年后个股动量效应减弱，故采用行业维度')
    add_text(doc, '3. 交易成本优势：ETF无印花税，成本更低')
    
    doc.add_heading('（三）风险管理必要性', 2)
    add_text(doc, '1. 动量崩盘风险：急速风格反转期存在左尾风险')
    add_text(doc, '2. 风控改善夏普比：止损、限仓能显著改善风险调整收益')
    add_text(doc, '3. 带宽再平衡：优于固定频率，减少不必要换手')
    
    doc.add_heading('（四）可转债优势', 2)
    add_text(doc, '1. 凸性特征：上行参与、下行缓冲')
    add_text(doc, '2. 适合小资金：20万级别资金可精细化操作')
    add_text(doc, '3. 中国特色：国内可转债市场活跃，机会丰富')
    
    doc.add_page_break()
    
    # ========== 预期收益评估 ==========
    doc.add_heading('六、预期收益评估', 1)
    
    doc.add_heading('（一）12个月收益预期', 2)
    
    return_table = doc.add_table(rows=4, cols=6)
    return_table.style = 'Table Grid'
    
    return_header = return_table.rows[0].cells
    return_header[0].text = '市场情景'
    return_header[1].text = '核心收益'
    return_header[2].text = '行业动量'
    return_header[3].text = '可转债'
    return_header[4].text = '综合收益'
    return_header[5].text = '预期金额'
    
    for cell in return_header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    return_data = [
        ['保守(熊市)', '-3%', '-20%', '-3%', '-6.4%', '18.72万'],
        ['基准(中性)', '+5%', '+15%', '+6%', '+7.1%', '21.42万'],
        ['乐观(牛市)', '+8%', '+35%', '+10%', '+13.6%', '22.72万']
    ]
    
    for i, row_data in enumerate(return_data, 1):
        row_cells = return_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    doc.add_heading('（二）盈亏场景分析', 2)
    
    add_text(doc, '容易盈利的市场环境：', bold=True)
    add_text(doc, '1. 行业分化明显，趋势连续', indent=1)
    add_text(doc, '2. 板块轮动有序，强弱分明', indent=1)
    add_text(doc, '3. 市场整体向上，动量效应显著', indent=1)
    
    add_text(doc, '容易亏损的市场环境：', bold=True)
    add_text(doc, '1. 横盘震荡，假突破频繁', indent=1)
    add_text(doc, '2. 政策扰动导致急涨急跌', indent=1)
    add_text(doc, '3. 行业相关性升高，同涨同跌', indent=1)
    add_text(doc, '4. 风格快速切换，动量失效', indent=1)
    
    doc.add_page_break()
    
    # ========== 中国市场本土化建议 ==========
    doc.add_heading('七、中国市场本土化建议', 1)
    
    doc.add_heading('（一）交易时间优化', 2)
    add_text(doc, '建议交易时间：9:30-10:00（开盘后稳定期）或14:30-15:00（尾盘）')
    add_text(doc, '避开时段：11:30和13:00前后的流动性低点')
    add_text(doc, '特殊时期：重要会议、节假日前后谨慎操作')
    
    doc.add_heading('（二）A股特色考虑', 2)
    add_text(doc, 'T+1交易制度：当日买入次日才能卖出，需预留流动性')
    add_text(doc, '涨跌停限制：主板10%、创业板/科创板20%，影响止损执行')
    add_text(doc, 'ST股票处理：避免ST股票，关注退市风险')
    add_text(doc, '散户市场特征：情绪化明显，需要更严格风控')
    
    doc.add_heading('（三）税费成本优化', 2)
    add_text(doc, 'ETF交易：免印花税，仅有佣金（通常万2.5）')
    add_text(doc, '可转债优势：深市可转债免佣金，成本更低')
    add_text(doc, '分红税收：持有期超过1年免税，1个月-1年征10%，1个月内征20%')
    
    doc.add_heading('（四）政策风险关注', 2)
    add_text(doc, '监管政策：关注证监会、交易所政策变化')
    add_text(doc, '行业政策：特别关注教育、地产、互联网平台等敏感行业')
    add_text(doc, '货币政策：央行利率决议对市场影响较大')
    
    doc.add_page_break()
    
    # ========== 实施建议 ==========
    doc.add_heading('八、实施建议', 1)
    
    doc.add_heading('（一）适合人群', 2)
    add_text(doc, '1. 有一定投资经验的个人投资者')
    add_text(doc, '2. 能承受中等风险')
    add_text(doc, '3. 有时间进行月度调仓')
    add_text(doc, '4. 理解并认同动量投资理念')
    add_text(doc, '5. 资金量在20-50万区间')
    
    doc.add_heading('（二）执行要点', 2)
    add_text(doc, '1. 严格执行纪律：不随意更改规则，避免情绪化操作')
    add_text(doc, '2. 做好记录：详细记录每次交易和调仓原因')
    add_text(doc, '3. 定期复盘：月度总结策略执行情况和市场变化')
    add_text(doc, '4. 保持学习：持续关注市场变化和策略改进机会')
    add_text(doc, '5. 风险控制：始终把风险控制放在第一位')
    
    doc.add_heading('（三）工具准备', 2)
    add_text(doc, '1. 券商选择：选择佣金低、系统稳定的券商')
    add_text(doc, '2. 交易软件：熟练使用券商APP和PC客户端')
    add_text(doc, '3. 数据来源：Wind、同花顺、东方财富等获取行情数据')
    add_text(doc, '4. 记录工具：Excel或专业软件进行策略计算和记录')
    add_text(doc, '5. 学习资源：关注优质投资公众号和研报')
    
    doc.add_heading('（四）心理准备', 2)
    add_text(doc, '1. 接受波动：股市波动是常态，要有心理准备')
    add_text(doc, '2. 长期视角：不要期望一夜暴富，坚持长期投资')
    add_text(doc, '3. 独立思考：不盲从市场情绪和小道消息')
    add_text(doc, '4. 止损纪律：达到止损线必须执行，不要心存侥幸')
    
    doc.add_page_break()
    
    # ========== 风险提示 ==========
    doc.add_heading('九、风险提示', 1)
    
    p = doc.add_paragraph()
    p.add_run('重要提示：').bold = True
    p.add_run('股市有风险，投资需谨慎。以下风险需要特别注意：').font.color.rgb = RGBColor(255, 0, 0)
    
    add_text(doc, '1. 市场风险：股市可能大幅波动，可能面临本金损失')
    add_text(doc, '2. 策略风险：动量策略存在失效风险，特别是市场风格切换时')
    add_text(doc, '3. 流动性风险：极端市场情况下可能面临流动性不足')
    add_text(doc, '4. 操作风险：需要纪律执行，情绪化操作会破坏策略效果')
    add_text(doc, '5. 监管风险：政策变化可能影响策略执行')
    add_text(doc, '6. 技术风险：交易系统故障可能影响策略执行')
    add_text(doc, '7. 集中度风险：行业集中可能放大单一行业风险')
    
    doc.add_paragraph()
    
    # 免责声明
    disclaimer = doc.add_paragraph()
    disclaimer.add_run('免责声明：').bold = True
    disclaimer.add_run('本策略仅供参考学习，不构成任何投资建议。投资者应根据自身情况独立判断，风险自担。过往业绩不代表未来表现。')
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer.runs[1].font.size = Pt(10)
    disclaimer.runs[1].font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    doc.save('/Users/maokaiyue/momentum-lens/炒股策略_优化版.docx')
    print("文档已成功创建：炒股策略_优化版.docx（无项目符号版本）")

if __name__ == "__main__":
    create_strategy_document()